import glob
from logging import root
from tkinter import DISABLED
from flask import Flask, redirect, render_template, request, send_file, send_from_directory, current_app as app, flash, url_for, session
from flask_login import login_required

from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.datastructures import MultiDict

from pymongo import MongoClient
from flask_mongoengine import MongoEngine
from flask_wtf import FlaskForm
from wtforms import SelectField, IntegerField, StringField, SubmitField, PasswordField, EmailField
from wtforms import MultipleFileField
from wtforms.validators import DataRequired, Email, EqualTo, Length
import pandas as pd
from datetime import datetime, timedelta
from docx2pdf import convert
from docx import Document
import os
import sys
import pythoncom
import traceback

# Gdrive Helper functions
#from gdrive_upload import find_folder, create_folder, callback, grant_permission, upload_to_gdrive
from gdrive_upload import *

# Routes
from functools import wraps

# Setup Flask app
application = Flask(__name__)
application.config.from_object('config.DevConfig')

# Connect to Mongodb
db = MongoEngine(application)

client = MongoClient(application.config['MONGO_URI'])
mongo_db = client.TEST
addresses = pd.DataFrame(mongo_db.get_collection(
    'Address').find({}, {"_id": 0, 'address': 1}))
adds = addresses.address.tolist()

# Connect to tables
ips = pd.DataFrame(mongo_db.get_collection(
    'ref_config').find({"Key": "InfluencerProgram"}, {"_id": 0, 'Values': 1}))
channels = pd.DataFrame(mongo_db.get_collection(
    'ref_config').find({"Key": "Channel"}, {"_id": 0, 'Values': 1}))

users_collection = mongo_db['users']
user_profile_collection = mongo_db['user_profile']


# Decorators


def login_required(f):
    @wraps(f)
    def wrap(*args, **kwards):
        if 'email' in session:
            return f(*args, **kwards)
        else:
            return redirect('/')

    return wrap


############ HELPER FUNCTIONS ############
def generate_invoice(data_df):
    e = ""
    try:
        print(os.getcwd(), file=sys.stderr)
        document = Document(os.path.join(os.getcwd(), 'invoice_template.docx'))
        data = {
            "[Name]": "",
            "[address1]": "",
            "[address2]": "",
            "[address3]": "",
            "[Invoice#]": "",
            "[Date]": "",
            "[caddress1]": "",
            "[caddress2]": "",
            "[caddress3]": "",
            "[caddress4]": "",
            "[desc1]": '',
            "[amount1]": '',
            "[desc2]": '',
            "[amount2]": '',
            "[desc3]": '',
            "[amount3]": '',
            "[TotalAmount]": "",
            "[bankName]": '',
            "[acctHolder]": "",
            "[acctNumber]": 0,
            "[IFSC]": '',
            "[PAN]": ''
        }

        # Update data dict
        dict_keys = {k.lower() for k in data.keys()}
        for col in data_df.columns:
            if "[" + col.lower() + "]" in dict_keys:
                data["[" + col + "]"] = data_df[col].values[0]

        for i, val in enumerate(data_df['BillTO'].values[0].split("|")):
            data['[caddress{}]'.format(i+1)] = val.strip()

        for idx, val in enumerate(data_df['Invoice-details'][0]):
            data['[desc{}]'.format(idx+1)] = val['desc']
            data['[amount{}]'.format(idx+1)] = val['amount']

        data['[address1]'] = data_df['Address'].values[0]
        # Update invoice
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip() in data.keys():
                            paragraph.text = str(
                                data[paragraph.text.strip()])

        # Create invoice docx
        user_path = session['path']
        filename = session['month']+session['program'] + \
            session['pan']+data['[Invoice#]']
        print(user_path)
        invoice_name = os.path.join(
            user_path, 'invoice_{}.docx'.format(filename))
        document.save(invoice_name)
        # Create invoice pdf
        pdf_name = os.path.join(
            user_path, '{}.pdf'.format(filename))
        pythoncom.CoInitialize()
        convert(invoice_name, pdf_name)
        # Delete docx
        os.remove(invoice_name)
    except:
        print(traceback.format_exc(), file=sys.stderr)
        return "error"

    return "success", '{}.pdf'.format(filename)


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower(
           ) in application.config['ALLOWED_EXTENSIONS']


def upload_files(request, upload_button, file_type):
    # check if the post request has the file part
    '''if upload_button not in request.files:
        print('No file part')
        return redirect(request.url)'''
    file = request.files[upload_button]
    print(request.files.getlist(upload_button))
    # If the user does not select a file, the browser submits an
    # empty file without a filename.

    if file.filename == '':
        print('No selected file')
        return redirect(request.url)
    if file and allowed_file(file.filename):
        for f in request.files.getlist(upload_button):
            filename = secure_filename(f.filename)
            if '.pdf' in filename:
                f.save(os.path.join(
                    application.config['UPLOAD_FOLDER'], session['email'], filename))
        flash('{} files uploaded'.format(file_type), "upload")
    else:
        flash("Only pdf files are allowed", "upload")
        # return redirect(url_for('upload_file', name=filename))


def setup_dir():
    base_dir = os.path.join(os.getcwd(), 'static/uploads')
    user_dir = os.path.join(base_dir, session['email'])

    if not os.path.exists(user_dir):
        os.makedirs(user_dir)
    else:
        files = glob.glob(user_dir+"/*")
        for f in files:
            os.remove(f)
    session['path'] = user_dir


class InvoiceForm(FlaskForm):
    name = StringField('Name', validators=[DataRequired()])
    address = StringField('Address', validators=[DataRequired()])
    billto = SelectField('Bill To', validate_choice=True, choices=[])
    infProgram = SelectField('Influencer Program',
                             validate_choice=True, choices=[])
    campaignMonth = SelectField('Month', validate_choice=True, choices=[])
    channel = SelectField('Channel', validate_choice=True, choices=[])
    # Billing Description
    description1 = StringField('Description', validators=[
                               DataRequired()], render_kw={'readonly': True})
    amount1 = IntegerField('Amount', validators=[DataRequired()])
    file1 = MultipleFileField(render_kw={'multiple': True})
    description2 = StringField('Description', render_kw={'readonly': True})
    amount2 = IntegerField('Amount')
    file2 = MultipleFileField(render_kw={'multiple': True})
    description3 = StringField('Description', render_kw={'readonly': True})
    amount3 = IntegerField('Amount')
    file3 = MultipleFileField(render_kw={'multiple': True})
    # Bank Details
    bankName = StringField('Bank Name', validators=[DataRequired()])
    acctHolder = StringField('Account Holder Name',
                             validators=[DataRequired()])
    acctNumber = StringField('Account Number', validators=[DataRequired()])
    IFSC = StringField('IFSC', validators=[DataRequired()])
    PAN = StringField('PAN', validators=[DataRequired()])

    # remember_me = BooleanField('Remember Me')
    submit = SubmitField('Submit')


class LoginForm(FlaskForm):
    username = StringField('UserName', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired()])
    signin = SubmitField('SignIn')
    test = StringField('test')


class RegistrationForm(FlaskForm):
    firstname = StringField('Firstname', validators=[DataRequired()])
    lastname = StringField('Lastname')
    email = EmailField('Email', validators=[DataRequired(), Email()])
    password1 = PasswordField('Password', validators=[DataRequired()])
    password2 = PasswordField('Confirm Password', validators=[
                              DataRequired(), EqualTo('password1')])
    phonenum = IntegerField('Phone Number', validators=[
                            Length(min=10, max=10)])
    # Bank Details
    bankName = StringField('Bank Name', validators=[DataRequired()])
    acctHolder = StringField('Account Holder Name',
                             validators=[DataRequired()])
    acctNumber = StringField('Account Number', validators=[DataRequired()])
    IFSC = StringField('IFSC', validators=[DataRequired()])
    PAN = StringField('PAN', validators=[DataRequired()])
    submit = SubmitField()


@application.route('/register1', methods=['POST', 'GET'])
def register():
    form = RegistrationForm()
    # if form.validate_on_submit():
    if request.method == "POST":
        firstname = form.firstname.data
        lastname = form.lastname.data
        email = form.email.data
        password = form.password1.data
        cpassword = form.password2.data
        bankname = form.bankName.data
        acctNumber = form.acctNumber.data
        acctHolder = form.acctHolder.data
        ifsc = form.IFSC.data
        pan = form.PAN.data

        userexists = list(users_collection.find({"user": email}, {"_id": 1}))
        if password != cpassword:
            flash('Passwords do not match')
        elif len(userexists) > 0:
            flash("Email already registered, try logging in.")
        else:
            update = users_collection.insert_one(
                {"user": email, "password": generate_password_hash(password)})
            bankdetails = {"bankname": bankname,
                           "acctNumber": acctNumber, "acctHolder": acctHolder, "IFSC": ifsc, "PAN": pan}
            update_up = user_profile_collection.insert_one(
                {"email": email, "firstname": firstname, "lastname": lastname, "bankDetails": bankdetails})
            session['email'] = email
            # Create folder
            folder_name = firstname + pan
            create_folder(folder_name)
            return redirect(url_for('invoice'))
    return render_template('register.html', form=form)


@application.route('/register', methods=['POST', 'GET'])
def test():
    form = RegistrationForm()
    if request.method == "POST":
        firstname = form.firstname.data
        lastname = form.lastname.data
        email = form.email.data
        password = form.password1.data
        cpassword = form.password2.data
        bankname = form.bankName.data
        acctNumber = form.acctNumber.data
        acctHolder = form.acctHolder.data
        ifsc = form.IFSC.data
        pan = form.PAN.data

        userexists = list(users_collection.find({"user": email}, {"_id": 1}))
        if password != cpassword:
            flash('Passwords do not match')
        elif len(userexists) > 0:
            flash("Email already registered, try logging in.")
        else:
            update = users_collection.insert_one(
                {"user": email, "password": generate_password_hash(password)})
            bankdetails = {"bankname": bankname,
                           "acctNumber": acctNumber, "acctHolder": acctHolder, "IFSC": ifsc, "PAN": pan}
            update_up = user_profile_collection.insert_one(
                {"email": email, "firstname": firstname, "lastname": lastname, "bankDetails": bankdetails})
            session['email'] = email
            # Create folder
            folder_name = firstname + pan
            create_folder(folder_name)
            return redirect(url_for('invoice'))
    return render_template('signup.html', form=form)


@application.route('/', methods=['GET', 'POST'])
@application.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    loginform = LoginForm()
    username = loginform.username.data
    password = loginform.password.data

    message = "Please login to your account."
    flash((message))
    if request.method == "POST":
        if request.form['action'] == 'Sign In':
            ret = list(users_collection.find(
                {"user": username}, {"_id": 0, "password": 1}))
            print(len(ret))
            if len(ret) > 0:
                if check_password_hash(ret[0]['password'], password):
                    session['email'] = username
                    setup_dir()
                    return redirect(url_for('invoice'))
                else:
                    message = "Wrong password!"
                    flash("Wrong password!")
                    # return redirect(url_for('login'))
                    return render_template('login.html', form=loginform, error=error, message=message)
            else:
                message = "E-mail not found! Please signup."
                flash(message)
                # return redirect(url_for('login'))
                return render_template('login.html', form=loginform, error=error, message=message)
        else:
            return redirect(url_for('register'))

    return render_template('login.html', form=loginform, error=error, message=message)


@application.route("/logout", methods=["POST", "GET"])
def logout():
    if "email" in session:
        session.pop("email", None)
        loginform = LoginForm()
        error = None
        message = "You have been logged out!"
        return redirect(url_for('login'))
    else:
        return redirect(url_for('login'))


@ application.route('/display', methods=['GET', 'POST'])
def display_invoice():
    # return send_from_directory(application.config['UPLOAD_FOLDER'], '20220605174629.pdf')
    if request.method == "POST":
        if request.form['action'] == 'Submit':

            # Move files to gdrive
            foldername = session['name'] + session['pan']
            print(foldername)
            root_folder_id = find_folder(foldername)

            if root_folder_id == "":
                root_folder_id = create_folder(
                    foldername, parent_id=['1u8itRD21yYt_JXrFiIv_G0QqRKsgDfod'])

            # create first sub folder
            subfolder_1 = session['program']+session['pan']
            print(subfolder_1)
            parent_folder_1 = find_folder(subfolder_1)
            if parent_folder_1 == "":
                print("sub folder 1 ", subfolder_1, root_folder_id)
                parent_folder_1 = create_folder(
                    subfolder_1, parent_id=[root_folder_id])

            # create first sub folder
            subfolder_2 = session['month']+session['program']+session['pan']
            parent_folder_2 = find_folder(subfolder_2)
            if parent_folder_2 == "":
                parent_folder_2 = create_folder(
                    subfolder_2, parent_id=[parent_folder_1])
            print(root_folder_id, parent_folder_1, parent_folder_2)

            # Copy the files to the directory
            path = session['path']
            print(session['path'])
            for file in glob.glob(path+"/*.pdf"):
                print(file, parent_folder_2)
                upload_to_gdrive(file.split("\\")[-1], file, parent_folder_2)

            return redirect(url_for('invoice'))

    return render_template('display_invoice.html')


@ application.route("/invoice", methods=['GET', 'POST'])
@login_required
def invoice():
    form = InvoiceForm()
    flash('')
    form.billto.choices = adds
    form.infProgram.choices = ips['Values'].tolist()[0]
    form.channel.choices = channels['Values'].tolist()[0]
    now = datetime.now()
    result = [now.strftime("%B")]
    for _ in range(0, 3):
        now = now.replace(day=1) - timedelta(days=1)
        result.append(now.strftime("%B"))
    form.campaignMonth.choices = result

    # Fill in the user details
    user_data = pd.DataFrame(mongo_db.get_collection('user_profile').find(
        {"email": session['email']}))
    form.name.data = user_data['firstname'][0] + \
        " " + user_data['lastname'][0]
    form.bankName.data = user_data['bankDetails'][0]['bankname']
    form.acctHolder.data = user_data['bankDetails'][0]['acctHolder']
    form.acctNumber.data = user_data['bankDetails'][0]['acctNumber']
    form.IFSC.data = user_data['bankDetails'][0]['IFSC']
    form.PAN.data = user_data['bankDetails'][0]['PAN']

    if request.method == 'GET':
        formdata = session.get('formdata', None)
        if formdata:
            flash('')
            # print(formdata)
            form = InvoiceForm(MultiDict(formdata))
            form.infProgram.choices = ips['Values'].tolist()[0]
            form.channel.choices = channels['Values'].tolist()[0]
            form.campaignMonth.choices = result
            form.billto.choices = adds
            if formdata['amount1'] == "":
                form.amount1.data = 0
            if formdata['amount2'] == "":
                form.amount2.data = 0
            if formdata['amount3'] == "":
                form.amount3.data = 0
            #form = form(CombinedMultiDict((request.files, request.form)))
            form.validate()
            session.pop('formdata')
        return render_template('index.html', form=form)

    if request.method == "POST":
        if request.form['action'] == 'Submit':
            # if form.submit.data:
            invoice_list = []
            name = form.name.data
            address = form.address.data
            billto = form.billto.data
            infProg = form.infProgram.data
            channel = form.channel.data
            month = form.campaignMonth.data
            #####
            desc1 = form.description1.data
            amt1 = form.amount1.data
            desc2 = form.description2.data
            amt2 = form.amount2.data
            desc3 = form.description3.data
            amt3 = form.amount3.data
            invoice_num = datetime.today().strftime('%Y%m%d%H%M%S')
            date = datetime.today().strftime('%Y-%m-%d')
            bankName = form.bankName.data
            acctHolder = form.acctHolder.data
            acctNumber = form.acctNumber.data
            ifsc = form.IFSC.data
            PAN = form.PAN.data

            # Set session
            session['name'] = name
            session['pan'] = PAN
            session['program'] = infProg
            session['month'] = month

            # Get total amount
            totamt = 0.0
            for k in range(1, 4):
                try:
                    totamt += float(eval('amt{}'.format(k)))
                except:
                    pass

            # Generate billing description list
            for i in range(1, 4):
                invoice_list.append(
                    {"desc": eval('desc{}'.format(i)), "amount": eval('amt{}'.format(i))})

            # Prepare df to write to Mongo collection
            data_df = pd.DataFrame([[name, address, billto, infProg, channel, month, invoice_list, invoice_num, date, totamt, bankName, acctHolder, acctNumber, ifsc, PAN]],
                                   columns=['Name', 'Address', 'BillTO', 'Influencer Program', 'Channel', 'Month', 'Invoice-details', 'Invoice#', 'Date', 'TotalAmount', 'bankName', 'acctHolder', 'acctNumber', 'IFSC', 'PAN'])

            # Insert data into collection
            insert_data = mongo_db.get_collection('INVOICE_DETAILS').insert_many(
                data_df.to_dict('records'))

            session['formdata'] = request.form
            # Delete old files
            files = glob.glob(session['path']+"/*")
            files_to_delete = [
                x for x in files if PAN and datetime.now().strftime("%B") in x]
            for f in files_to_delete:
                os.remove(f)
            # Generate Invoice
            err, filename = generate_invoice(data_df)
            session['invoice'] = filename
            if err == "success":
                print("Invoice generated successfully", filename)
                return redirect(url_for('display_invoice'))
                # return render_template('display_invoice.html', path=session['email'], doc_id=filename)
            else:
                return "Error occured: {}".format(err)

        elif request.form['action'] == 'Upload PR':
            #print("two ", request.form['action'])
            # print(request.files['file1'])
            upload_files(request, 'file1', 'Product reimbursement')

        elif request.form['action'] == 'Upload TE':
            #print("three ", request.form['action'])
            upload_files(request, 'file3', 'Travel Expense')

    return render_template('index.html', form=form)


if __name__ == "__main__":
    application.run(debug=True)
